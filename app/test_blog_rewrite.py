import sys
import os
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))



import os
from docx import Document
from helper.docx_utils import extract_text_from_docx
from ui.components.query_handler import search_and_generate_response
from io import BytesIO
from datetime import datetime

# Load the input test blog doc
input_path = "data/client_samples/BNBGPT-Test-Blog-Posts-2.docx"
raw_text = extract_text_from_docx(input_path)

# Optional: If multiple blogs are separated by lines or markers
# For now, assume it's one long blog
blog_posts = [raw_text.strip()]

# Output doc
output_doc = Document()
output_doc.add_heading("BNBGPT – Rewritten Test Blogs", 0)
output_doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
output_doc.add_paragraph("")

# Run each blog through RAG
for idx, blog in enumerate(blog_posts, 1):
    result = search_and_generate_response(blog, k=5)

    output_doc.add_heading(f"🔹 Rewritten Blog Post #{idx}", level=1)
    output_doc.add_paragraph(result["response"])
    output_doc.add_paragraph("")

# Save to file
output_path = "data/client_samples/BNBGPT-Rewritten-Test-Blogs-2.docx"
output_doc.save(output_path)
print(f"✅ Rewritten blogs saved to: {output_path}")
print(f"the rewritten blog is {output_doc}")
